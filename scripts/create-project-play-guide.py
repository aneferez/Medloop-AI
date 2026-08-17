from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "MedLoop_AI_Project_and_Google_Play_Guide.docx"
LOGO = ROOT / "public" / "medloop-logo-512.png"

NAVY = "16324F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "166534"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend([run_color, underline])
    run.append(run_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for font_node in (qn("w:ascii"), qn("w:hAnsi"), qn("w:eastAsia")):
        normal.element.rPr.rFonts.set(font_node, "Calibri")
    settings = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in settings.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_para(doc, text="", bold_prefix=None, italic=False, color=None, keep=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        run = p.add_run(text)
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.keep_together = keep
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_steps(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, indent])
    level.extend([start, num_fmt, level_text, suffix, p_pr])
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.append(num_pr)
        p.add_run(item)


def add_callout(doc, label, text, kind="info"):
    colors = {"info": (LIGHT_BLUE, DARK_BLUE), "ready": ("E8F5E9", GREEN), "warning": ("FFF8E1", GOLD), "risk": ("FDECEC", RED)}
    fill, accent = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(label.upper() + "  ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(accent)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, heading in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(heading)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(value))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_source(doc, title, url, note):
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, title, url)
    p.add_run(f" - {note}")


def add_page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "MEDLOOP AI  |  PROJECT & GOOGLE PLAY RELEASE GUIDE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(MID_GRAY)
    add_page_number(section.footer.paragraphs[0])
    section.footer.paragraphs[0].runs[0].font.size = Pt(9)
    section.footer.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(MID_GRAY)

    # Cover: report-cover archetype using restrained MedLoop branding.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(1.25))
    p.paragraph_format.space_after = Pt(24)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("PRODUCT REFERENCE & RELEASE OPERATIONS")
    kr.bold = True
    kr.font.size = Pt(10)
    kr.font.color.rgb = RGBColor.from_string(BLUE)
    kicker.paragraph_format.space_after = Pt(12)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("MedLoop AI")
    tr.bold = True
    tr.font.size = Pt(30)
    tr.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(7)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Complete Project Details and Google Play Publishing Guide")
    sr.font.size = Pt(15)
    sr.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(24)
    add_callout(doc, "Purpose", "A single operational reference covering the product, architecture, privacy and safety model, build artifacts, release readiness, Play Console declarations, testing tracks, and production rollout.", "info")
    meta = add_table(doc, ["Document status", "Project snapshot"], [
        ("Prepared", "9 August 2026"),
        ("App version", "1.1.0-beta.11 (version code 15)"),
        ("Android package", "com.medloop.ai"),
        ("Primary artifact", "Signed Android App Bundle (.aab)"),
        ("Owner note", "Developed by Aneruth | Rosaline"),
    ], [2600, 6760])
    add_para(doc, "This guide is a technical and operational reference, not legal or regulatory advice. Google Play policies and applicable health/privacy laws must be reviewed again immediately before submission.", italic=True, color=MID_GRAY)

    add_page_break(doc)
    doc.add_heading("Document control and quick status", level=1)
    add_table(doc, ["Area", "Current status", "Release implication"], [
        ("Application", "Builds and launches on Android", "Ready for internal testing"),
        ("Automated QA", "23 tests, lint, and production build passed", "Retest after every release change"),
        ("Target SDK", "API 36", "Meets the announced 31 Aug 2026 requirement"),
        ("Signing", "RSA 4096 upload certificate; APK signature verified", "Keep keystore and credentials private"),
        ("Store listing", "Not completed in Play Console", "Requires text, graphics, support contact, privacy URLs"),
        ("Health policy", "In-app disclaimer exists", "Health declaration and store disclaimer still required"),
        ("Multilingual AI voice", "English, Hindi, Tamil Home MP3s bundled", "Other section/language recordings remain fallback TTS"),
        ("Production access", "Unknown external account state", "New personal accounts may require 12 testers / 14 days"),
    ], [1800, 2800, 4760])
    add_callout(doc, "Current recommendation", "Upload the AAB to Internal testing first. Do not start Production until the public privacy and deletion URLs, health declaration, Data safety answers, store assets, exact-alarm behavior, and closed-test requirements are complete.", "warning")

    doc.add_heading("Contents", level=1)
    add_bullets(doc, [
        "Part I - Product purpose, users, capabilities, and user journey",
        "Part II - Architecture, storage, privacy, security, Android configuration, and QA",
        "Part III - Release artifacts and Google Play readiness assessment",
        "Part IV - Step-by-step Google Play Console publishing procedure",
        "Part V - Proposed store listing copy, release checklist, and official sources",
    ])

    add_page_break(doc)
    doc.add_heading("Part I - Product overview", level=1)
    doc.add_heading("1. Purpose and positioning", level=2)
    add_para(doc, "MedLoop AI is a local-first Android medicine-coordination application. It helps an individual or family organize medicine schedules, record dose events, track stock, retain prescription and appointment information, prepare family-message drafts, and keep an emergency summary available on the device.")
    add_callout(doc, "Medical boundary", "MedLoop AI is an organization and reminder tool. It is not a medical device, diagnostic system, treatment recommender, emergency-monitoring service, or clinical decision-support system.", "risk")
    doc.add_heading("2. Intended users", level=2)
    add_bullets(doc, [
        "Adults managing their own medication schedules and appointments.",
        "Adult family caregivers coordinating routine care for loved ones.",
        "Users who prefer local device storage instead of a provider-operated cloud account.",
        "Users who understand that reminders and recorded adherence do not replace instructions from a qualified healthcare professional.",
    ])
    doc.add_heading("3. Core value proposition", level=2)
    add_bullets(doc, [
        "Local-first privacy: no MedLoop backend, analytics, advertising, Firebase, or cloud database in the reviewed build.",
        "Practical daily coordination: schedules, reminders, Taken/Missed actions, stock alerts, appointments, reports, and emergency details.",
        "User-controlled sharing: encrypted backup export and user-reviewed SMS/WhatsApp drafts.",
        "Contextual assistance: section guidance, approved non-medical help search, voice playback, and anonymous on-device helpfulness feedback.",
    ])

    doc.add_heading("4. Functional areas", level=2)
    add_table(doc, ["Section", "What the user can do"], [
        ("Authentication", "Create a local account, sign in, persist a local session, sign out, or delete the account with confirmation."),
        ("Home", "See setup progress and shortcuts for completing the initial care plan."),
        ("Dashboard", "View the next dose, mark Taken or Missed, switch among Halo, Timeline, and Companion layouts, and review alerts/appointments."),
        ("Family", "Create care profiles, define one Level 1 contact, associate medicines, and prepare refill-message drafts."),
        ("Medicines", "Save medicine label details, dose periods/times, stock, refill state, units per dose, and stock buffer."),
        ("Prescriptions", "Save prescriber/clinic/notes and optional camera or gallery images in local storage."),
        ("Alerts", "Review derived missed-dose, refill, and low-stock notices."),
        ("Appointments", "Save, edit, and delete upcoming doctor or clinic visits."),
        ("Reports", "Review today’s completion and up to the latest 30 displayed events from the retained dose history."),
        ("Emergency Card", "Review saved profile, allergy, blood group, contact, and medicine information."),
        ("Settings", "Manage identity, reminders, message preferences, encrypted backups, restore, privacy, and account deletion."),
        ("Privacy & Safety", "Read local data handling, notification limitations, disclaimer, retention, and deletion information."),
    ], [1900, 7460])

    add_page_break(doc)
    doc.add_heading("5. Typical user journey", level=2)
    add_steps(doc, [
        "Create a local email/password account and sign in on the device.",
        "Add the person receiving care and, if needed, a primary Level 1 contact.",
        "Add medicine names and dosages exactly from the label or prescription, then choose clinician-prescribed dose periods and times.",
        "Enable Android notification access and test reminder delivery; grant exact-alarm access if the device requires it.",
        "Use Dashboard each day and record Taken or Missed only after verifying what occurred.",
        "Maintain appointments, prescriptions, stock levels, allergies, and emergency contact details.",
        "Export a password-encrypted backup and store the backup password separately.",
        "Use Settings to delete records or the complete local account when no longer required.",
    ])
    doc.add_heading("6. AI assistance and multilingual voice", level=2)
    add_bullets(doc, [
        "The assistant opens after sign-in and introduces each of 11 protected app sections.",
        "Search answers are selected from an approved local help knowledge base and do not provide diagnosis or treatment advice.",
        "Anonymous ‘Was this helpful?’ counts are saved only on the current device.",
        "English, Hindi, and Tamil interface guidance is implemented for the current assistant content.",
        "Generated ‘delicate’ voice MP3 recordings are currently bundled for the Home section in English, Hindi, and Tamil.",
        "When a generated MP3 is missing, the app falls back to Android or browser text-to-speech.",
        "An optional VITE_ASSISTANT_AUDIO_BASE_URL can point to a static online mirror; local MP3 and device speech remain fallbacks.",
    ])
    add_callout(doc, "Known gap", "The requested 22-language, 11-section prebuilt recording set is not complete. The current verified package contains 3 of the planned 242 Home/section language recordings. Do not advertise full 22-language prerecorded coverage in the Play listing.", "warning")

    add_page_break(doc)
    doc.add_heading("Part II - Technical and security details", level=1)
    doc.add_heading("7. Technology stack", level=2)
    add_table(doc, ["Layer", "Implementation"], [
        ("UI", "React 19, React DOM, Material UI, Emotion, Lucide React"),
        ("Build", "Vite 8, TypeScript toolchain, npm"),
        ("Android", "Capacitor 8, Gradle 8.14.3, JDK 21, Android SDK 36"),
        ("Native plugins", "App, Camera, Filesystem, Local Notifications, Share, Secure Storage, Text to Speech"),
        ("Persistence", "Capacitor Secure Storage for accounts/session/structured records; IndexedDB for images"),
        ("Testing", "Vitest, fake-indexeddb, Oxlint, Android physical-device smoke testing"),
    ], [1900, 7460])
    doc.add_heading("8. Architecture", level=2)
    add_para(doc, "React owns application state and domain behavior. Capacitor packages the Vite output into an Android WebView and bridges to Android services. There is no application API, server, cloud database, remote authentication provider, or project-owned background worker.")
    add_bullets(doc, [
        "src/App.jsx orchestrates authentication, persistence, CRUD, media, notifications, backup/restore, and navigation.",
        "src/components/AppShell.jsx supplies protected navigation and the signed-in layout.",
        "src/pages contains lazy-loaded section UI; src/lib contains domain and platform modules.",
        "Browser History API and internal maps provide navigation without a routing package.",
        "Capacitor copies dist into android/app/src/main/assets/public for release builds.",
    ])
    doc.add_heading("9. Storage and retention", level=2)
    add_table(doc, ["Data", "Location and control"], [
        ("Account/session", "Device-protected Capacitor Secure Storage; local-only authentication."),
        ("Care records", "Per-account structured state in Secure Storage."),
        ("Images", "IndexedDB within the application sandbox; JPEG/PNG/WebP up to 10 MB."),
        ("Dose history", "Newest 200 Taken/Missed events retained; reports show the latest relevant subset."),
        ("Backups", "User-exported .medloop file encrypted with PBKDF2-SHA-256 and AES-256-GCM."),
        ("Deletion", "In-app account deletion removes account, session, structured records, profile image, and owned prescription images."),
    ], [1900, 7460])

    doc.add_heading("10. Security and privacy model", level=2)
    add_bullets(doc, [
        "Passwords are not stored in plaintext; new verifiers use salted PBKDF2-HMAC-SHA-256 with 210,000 iterations.",
        "Encrypted backups use PBKDF2-HMAC-SHA-256 with 310,000 iterations and AES-256-GCM authenticated encryption.",
        "Android cloud backup and device-transfer extraction are disabled.",
        "Family messages are drafts opened in an external messaging app; MedLoop does not silently send them.",
        "No analytics, advertising, crash reporting, remote database, Firebase, or cloud AI SDK is active in the reviewed build.",
        "Notification content may expose medicine information on a lock screen; users should review device privacy settings.",
        "A rooted device, malicious accessibility/overlay software, or compromised operating system can defeat local safeguards.",
    ])
    add_callout(doc, "Compliance boundary", "Do not claim HIPAA, GDPR, India DPDP, medical-device, clinical, or other regulatory compliance without a dedicated legal, security, and operational assessment.", "risk")
    doc.add_heading("11. Android configuration", level=2)
    add_table(doc, ["Setting", "Verified value"], [
        ("App name", "MedLoop AI"),
        ("Application ID / namespace", "com.medloop.ai"),
        ("Version", "1.1.0-beta.11"),
        ("Version code", "15"),
        ("Minimum SDK", "24 (Android 7.0)"),
        ("Target / compile SDK", "36 (Android 16)"),
        ("Camera hardware", "Optional"),
        ("Android backup", "Disabled"),
        ("Release shrinking", "minifyEnabled false"),
    ], [2900, 6460])
    doc.add_heading("12. Effective Android permissions", level=2)
    add_table(doc, ["Permission", "Purpose / review note"], [
        ("INTERNET", "Capacitor WebView/runtime and optional static voice mirror; reviewed build has no MedLoop backend."),
        ("SCHEDULE_EXACT_ALARM", "User-granted special access for precise medicine reminders; not auto-granted on many Android 14+ installs."),
        ("POST_NOTIFICATIONS", "Runtime notification permission for medicine reminders."),
        ("RECEIVE_BOOT_COMPLETED", "Allows notification scheduling components to recover after restart."),
        ("WAKE_LOCK", "Supports reminder delivery by the notification plugin."),
        ("Dynamic receiver permission", "Application-scoped non-exported receiver protection generated by Android tooling."),
    ], [3000, 6360])
    add_callout(doc, "Exact alarm review", "The app declares SCHEDULE_EXACT_ALARM, not the more restricted USE_EXACT_ALARM. Verify the denied-permission flow on Android 14-16 and explain in the listing that precise reminder delivery depends on Android notification and exact-alarm settings.", "warning")

    add_page_break(doc)
    doc.add_heading("13. Quality assurance snapshot", level=2)
    add_table(doc, ["Check", "Latest verified result"], [
        ("Unit/integration tests", "23 passed across 10 test files"),
        ("Lint", "Oxlint passed"),
        ("Web production build", "Vite production build passed"),
        ("Android release", "Signed APK and AAB builds passed"),
        ("Signing inspection", "APK Signature Scheme v2; RSA 4096 signer"),
        ("Physical device", "Release APK installed and launched on connected Android device"),
        ("Bundled Home voices", "English 234,240 bytes; Hindi 312,192 bytes; Tamil 334,848 bytes"),
    ], [2800, 6560])
    doc.add_heading("14. Release artifacts", level=2)
    add_table(doc, ["Artifact", "Size", "SHA-256"], [
        ("MedLoop-AI-release.aab", "7,826,492 bytes", "C6CB530FABF2A1E008ECC110243364939E7EA09C662431FFFE8BEF2E86A139BD"),
        ("MedLoop-AI-release.apk", "8,780,261 bytes", "30B78BFB0DCFF8CF0A7BED087D86B89801C10C06794C64149F775C6FBDD2BFCE"),
    ], [2600, 1500, 5260])
    add_para(doc, "Use the AAB for Play Console. Use the APK only for direct device testing. Before any new upload, increase versionCode above every code already used in Play Console and rebuild both artifacts.")

    add_page_break(doc)
    doc.add_heading("Part III - Google Play readiness", level=1)
    doc.add_heading("15. Readiness assessment", level=2)
    add_table(doc, ["Requirement", "Status", "Required action"], [
        ("Signed AAB", "READY", "Upload artifacts/MedLoop-AI-release.aab to Internal testing."),
        ("Target API", "READY", "Keep target API 36 or higher for submissions on/after 31 Aug 2026."),
        ("Unique package", "READY / VERIFY", "Reserve/register com.medloop.ai in the owner’s verified Play account."),
        ("Version code", "READY FOR FIRST UPLOAD", "If code 15 is already used, increment and rebuild."),
        ("Privacy policy", "FILE EXISTS", "Publish public/privacy.html at an active public non-geofenced HTTPS URL; no PDF."),
        ("Account deletion page", "FILE EXISTS", "Publish public/account-deletion.html and provide its URL where Play requests it."),
        ("Support contact", "BLOCKED EXTERNALLY", "Add a monitored support email and preferably a support website."),
        ("Health declaration", "NOT SUBMITTED", "Declare Medication and Treatment Management; confirm not a medical device."),
        ("Data safety", "DRAFT NEEDED", "Complete from actual production behavior and every dependency."),
        ("Store disclaimer", "NOT YET PUBLISHED", "Include Google-required non-medical-device disclaimer in full description."),
        ("Store graphics", "PARTIAL", "Prepare 512 icon, 1024x500 feature graphic, and at least two accurate phone screenshots."),
        ("Closed testing", "ACCOUNT DEPENDENT", "New personal account after 13 Nov 2023: 12 opted-in testers for 14 continuous days."),
        ("Production access", "ACCOUNT DEPENDENT", "Apply after required test and answer production-readiness questions."),
        ("22-language prerecorded voice", "INCOMPLETE", "Do not advertise as complete; finish translations, native review, generation, and bundling first."),
    ], [2600, 1800, 4960])
    add_callout(doc, "Go / no-go", "GO for Internal testing after the listing and privacy URLs are prepared. NO-GO for Production until policy declarations, public support/privacy details, closed-testing requirements, and final device/reminder testing are complete.", "warning")

    doc.add_heading("Part IV - Google Play upload procedure", level=1)
    doc.add_heading("16. Create and verify the developer account", level=2)
    add_steps(doc, [
        "Sign in with the organization’s long-term Google Account and create a Play Console developer account.",
        "Accept the Developer Distribution Agreement and pay Google’s one-time US$25 registration fee.",
        "Complete identity and contact verification. Enable 2-Step Verification for the owner and every privileged user.",
        "For newly created personal accounts, complete device verification using the Play Console mobile app when prompted.",
        "Confirm the package-registration status for com.medloop.ai. Google states all Play package names must be registered by 30 September 2026.",
        "Invite additional release users only with least-privilege roles; do not share the account owner password.",
    ])
    doc.add_heading("17. Prepare final release inputs", level=2)
    add_steps(doc, [
        "Choose the public release version. For a production candidate, replace the beta version name if desired and increase versionCode from 15 to an unused higher integer.",
        "Run npm install, npm run lint, npm test, npm run build, and npm run android:release.",
        "Verify the AAB/APK package, version, target SDK, permissions, signing certificate, hashes, and physical-device launch.",
        "Host privacy.html and account-deletion.html on stable public HTTPS URLs. Confirm they work without login, region restriction, downloads, or editable-document access.",
        "Create a monitored support email and public support/security contact page.",
        "Prepare approved store text, icon, feature graphic, screenshots, release notes, and testing instructions.",
        "Retest Android 14-16 notification permission, exact-alarm access, restart/rescheduling, lock-screen behavior, backup/restore, wrong-password failure, and complete account deletion.",
    ])
    add_callout(doc, "Key protection", "The repository uses an upload keystore. Keep the .jks file and passwords outside source control, back them up securely, and enroll in Play App Signing so Google protects the app-signing key while your upload key remains resettable.", "risk")

    add_page_break(doc)
    doc.add_heading("18. Create the app in Play Console", level=2)
    add_steps(doc, [
        "Open Play Console and select Home > Create app.",
        "Set the default language, enter MedLoop AI, select App rather than Game, and choose Free unless the commercial model intentionally changes.",
        "Enter the monitored developer contact email.",
        "Accept the Developer Program Policies, export-law declaration, and Play App Signing terms.",
        "Create the app and use the Dashboard setup tasks as the submission checklist.",
    ])
    doc.add_heading("19. Complete the main store listing", level=2)
    add_bullets(doc, [
        "App name: maximum 30 characters.",
        "Short description: maximum 80 characters.",
        "Full description: maximum 4,000 characters; describe only released functionality.",
        "App icon: 32-bit PNG with alpha, 512 x 512 px, maximum 1,024 KB.",
        "Feature graphic: JPEG or 24-bit PNG without alpha, exactly 1,024 x 500 px.",
        "Phone screenshots: at least two; JPEG or 24-bit PNG without alpha; each side 320-3,840 px and the longest side no more than twice the shortest side.",
        "Category recommendation: Medical. Add accurate contact details and a privacy-policy URL.",
        "Do not claim diagnosis, treatment, emergency monitoring, guaranteed reminder delivery, regulatory certification, or complete 22-language prerecorded voice coverage.",
    ])

    add_page_break(doc)
    doc.add_heading("20. Complete App content declarations", level=2)
    doc.add_heading("20.1 Privacy policy and account deletion", level=3)
    add_bullets(doc, [
        "Provide the public privacy-policy HTTPS URL in Play Console and keep the same link/text accessible inside the app.",
        "The policy must describe local health data, images, notifications, message/share handoffs, encrypted backups, optional voice hosting, retention, and deletion.",
        "Provide the account-deletion URL if Play requests a web deletion resource. The in-app Settings flow already supports complete local account deletion.",
    ])
    doc.add_heading("20.2 Data safety - evidence-based draft", level=3)
    add_para(doc, "For the currently reviewed local-only build, the likely declaration is that the developer does not collect or share user data because account and health records are not transmitted off the device. This is a draft, not an automatic final answer.")
    add_bullets(doc, [
        "Reconfirm every production SDK, network request, hosted audio URL, support form, and future service before submission.",
        "User-initiated export through the Android share sheet and user-reviewed message drafts must be accurately described in the privacy policy even when they are not developer collection.",
        "State that users can delete local data and accounts in the app.",
        "Do not claim an independent security review unless one has actually occurred.",
        "The developer account owner is responsible for the final declaration and any later updates when behavior changes.",
    ])
    doc.add_heading("20.3 Health apps declaration", level=3)
    add_bullets(doc, [
        "Select Medical > Medication and Treatment Management.",
        "Do not select Medical Device Apps unless the owner has applicable regulatory status and evidence.",
        "Disclose that the app organizes medication schedules, reminders, adherence records, prescription records, and related care information.",
        "Use the exact store disclaimer: ‘MedLoop AI is not a medical device and does not diagnose, treat, cure, or prevent any medical condition.’",
        "Keep the in-app warnings to follow the prescription/clinician and contact local emergency services in an emergency.",
    ])

    add_page_break(doc)
    doc.add_heading("20.4 Remaining declarations", level=3)
    add_table(doc, ["Declaration", "Recommended response / action"], [
        ("Ads", "No, based on the reviewed build. Revisit if any advertising SDK is added."),
        ("App access", "Explain that users can create a local account directly; provide review instructions if Google cannot reach protected screens."),
        ("Content rating", "Complete the IARC questionnaire accurately; update it whenever content changes."),
        ("Target audience", "Recommend adults 18+ unless the owner intentionally targets minors and completes all additional Families obligations."),
        ("News / government", "No, based on current functionality and ownership."),
        ("Financial features", "No, based on current functionality."),
        ("Permissions", "Review any Play-generated declaration alert after the AAB is uploaded; justify only permissions actually required."),
        ("Exact alarm", "The build uses user-granted SCHEDULE_EXACT_ALARM. Document reminder use and retain a graceful non-exact fallback."),
    ], [2600, 6760])

    doc.add_heading("21. Configure Play App Signing", level=2)
    add_steps(doc, [
        "When preparing the first release, accept Play App Signing and use Google’s recommended generated app-signing key unless cross-store signing requires another deliberate strategy.",
        "Upload the locally signed AAB using the project upload key.",
        "Record the Play app-signing certificate SHA-1/SHA-256 fingerprints and the local upload certificate separately.",
        "If future APIs authenticate by Android certificate, register the Google Play app-signing certificate fingerprint, not only the upload certificate.",
        "Keep the upload keystore and passwords in secured, backed-up storage; never commit them.",
    ])

    add_page_break(doc)
    doc.add_heading("22. Upload to Internal testing", level=2)
    add_steps(doc, [
        "Go to Test and release > Testing > Internal testing and create a release.",
        "Upload artifacts/MedLoop-AI-release.aab. Play Console will inspect package com.medloop.ai, version code 15, target API 36, permissions, and signing.",
        "Resolve every blocking error. If version code 15 was previously uploaded, increase versionCode and rebuild; uploaded version codes cannot be reused.",
        "Enter concise release notes, save the draft, review the release, and roll it out to the internal track.",
        "Add trusted tester email addresses or a Google Group and share the opt-in link.",
        "Install from Google Play, not by sideloading, then run the full tester checklist and inspect the pre-launch report, Android vitals, policy status, and testing feedback.",
    ])
    doc.add_heading("23. Closed testing and production access", level=2)
    add_para(doc, "For a personal developer account created after 13 November 2023, Google currently requires a closed test with at least 12 testers continuously opted in for the last 14 days before applying for Production access.")
    add_steps(doc, [
        "Finish required app setup and create a Closed testing track.",
        "Add at least 12 representative testers and ensure they opt in and remain opted in continuously for at least 14 days.",
        "Ask testers to install from Play, exercise authentication, medicine CRUD, reminders, dose logging, stock, prescriptions/camera, backup/restore, assistant, privacy, and deletion, and submit feedback.",
        "Keep a dated test plan, device/Android-version matrix, issues, fixes, and feedback summary.",
        "When Play Console enables the option, apply for Production access from the Dashboard and answer the testing, feedback, app-purpose, and production-readiness questions truthfully.",
    ])

    add_page_break(doc)
    doc.add_heading("24. Production release and rollout", level=2)
    add_steps(doc, [
        "After Production access is approved, go to Test and release > Production and create a new release or promote the tested bundle where available.",
        "Confirm countries/regions, pricing, device availability, store listing, support details, declarations, app-signing status, and release notes.",
        "Review Play Console warnings and the publishing overview. Do not send changes for review until all intended changes are included.",
        "Submit the release for review. Review time is controlled by Google and must not be represented as immediate or guaranteed.",
        "Use managed/staged publishing when timing matters. For the first public release, consider a controlled country or percentage rollout if Play makes that option available.",
        "Monitor crashes/ANRs, ratings, policy status, reminder complaints, privacy contacts, and deletion requests. Pause rollout if a safety or data issue appears.",
    ])
    doc.add_heading("25. Updating the app later", level=2)
    add_steps(doc, [
        "Increase versionCode for every upload and update versionName for the user-facing release.",
        "Update package.json, Android Gradle configuration, filenames/scripts, release notes, and this guide so versions remain aligned.",
        "Rerun tests, lint, production build, release build, signing inspection, and physical-device smoke tests.",
        "Update Data safety, Health, permissions, target audience, content rating, and privacy disclosures whenever functionality or SDK behavior changes.",
        "Upload the new AAB to Internal/Closed testing, validate Play-generated APKs, then promote to Production.",
    ])

    add_page_break(doc)
    doc.add_heading("Part V - Ready-to-use listing material", level=1)
    doc.add_heading("26. Proposed Play Store listing", level=2)
    add_para(doc, "App name (10/30 characters)", bold_prefix="App name")
    add_callout(doc, "Copy", "MedLoop AI", "info")
    add_para(doc, "Short description (under 80 characters)", bold_prefix="Short description")
    add_callout(doc, "Copy", "Private medicine reminders, care schedules and family coordination on device.", "info")
    add_para(doc, "Suggested full description", bold_prefix="Suggested full description")
    for paragraph in [
        "MedLoop AI helps adults and family caregivers organize medicine schedules, reminders, dose records, appointments, prescriptions, stock levels and emergency details on an Android device.",
        "Plan daily care: add medicine label details, select clinician-prescribed dose periods and times, and review the next scheduled dose on the dashboard. Record Taken or Missed actions only after confirming what happened.",
        "Coordinate locally: maintain family care profiles, appointments, prescription notes and images, refill information, and user-reviewed SMS or WhatsApp message drafts. MedLoop AI does not silently send messages.",
        "Protect your records: the reviewed release stores care records locally, supports in-app account deletion, disables Android cloud backup, and provides password-encrypted backup export and restore. Keep exported files and backup passwords secure and separate.",
        "Get contextual help: the built-in guide explains each section, searches approved non-medical help topics, provides English, Hindi and Tamil guidance, and uses generated or device voice where available.",
        "Reminder delivery depends on Android notification, battery and exact-alarm settings. Reports reflect only actions recorded in the app and are not proof that a medicine was or was not taken.",
        "MedLoop AI is not a medical device and does not diagnose, treat, cure, or prevent any medical condition. Follow the medicine label and instructions from a qualified healthcare professional. For urgent symptoms or an emergency, contact local emergency services immediately.",
    ]:
        add_para(doc, paragraph)
    doc.add_heading("27. Suggested release notes", level=2)
    add_bullets(doc, [
        "Added post-login contextual guidance for every app section.",
        "Added approved local help search and anonymous on-device helpfulness feedback.",
        "Added English, Hindi and Tamil Home voice guidance with offline playback and device-speech fallback.",
        "Improved Android reminder support, privacy documentation, encrypted backups, and release verification.",
    ])

    doc.add_heading("28. Store asset checklist", level=2)
    add_table(doc, ["Asset", "Specification", "Project status"], [
        ("App icon", "512x512, 32-bit PNG with alpha, <=1,024 KB", "public/medloop-logo-512.png exists; visually confirm Play icon safe zone."),
        ("Feature graphic", "1,024x500 JPEG or 24-bit PNG, no alpha", "Not identified; create before listing submission."),
        ("Phone screenshots", "At least 2, 320-3,840 px, valid aspect ratio", "Device screenshots exist in artifacts; select current, privacy-safe final images."),
        ("Tablet screenshots", "Recommended if tablet distribution is enabled", "Not verified."),
        ("Preview video", "Optional YouTube URL", "Not required."),
        ("Alt text", "Short accurate descriptions for visual assets", "Write during asset upload."),
    ], [2000, 3600, 3760])
    doc.add_heading("29. Tester acceptance checklist", level=2)
    add_bullets(doc, [
        "Fresh install, local sign-up, sign-in persistence, sign-out, and wrong-password handling.",
        "Create/edit/delete family member, medicine, prescription image, appointment, and alert-related records.",
        "Notification permission, exact-alarm access, 10-second test, scheduled delivery, reboot recovery, and denied-permission behavior.",
        "Taken/Missed transitions, stock deduction/restoration, low-stock alert, daily reset, and report accuracy.",
        "Assistant section navigation, English/Hindi/Tamil selection, Home MP3 playback, fallback TTS, help search, and feedback.",
        "Encrypted backup export, wrong-password rejection, restore confirmation, and restored media/records.",
        "SMS/WhatsApp drafts require user review and do not send automatically.",
        "Privacy/disclaimer visibility, complete account deletion, reinstall behavior, and absence of real-user data in screenshots/logs.",
        "Small/large text, screen reader labels, portrait layout, offline mode, low battery, and Android 7 through 16 compatibility as devices permit.",
    ])

    doc.add_heading("30. Final submission checklist", level=2)
    add_bullets(doc, [
        "Developer identity, contact information, 2-Step Verification, and package registration complete.",
        "Public privacy policy and account-deletion URLs live and checked without authentication.",
        "Support email monitored; security/privacy escalation route documented.",
        "Final versionCode unused and higher than every prior Play upload.",
        "AAB rebuilt, signed, hashed, and tested from the Play track.",
        "Store listing copy and screenshots accurately match the submitted build.",
        "Data safety, Health apps, ads, app access, target audience, content rating, and permissions declarations complete and consistent.",
        "Required disclaimer present in store description and in app.",
        "Closed-test requirement completed where applicable, with 12 opted-in testers for 14 continuous days.",
        "Pre-launch report, Android vitals, policy status, and tester feedback reviewed; critical issues fixed.",
        "Production release reviewed through Publishing overview and submitted by an authorized owner/release manager.",
        "Post-launch monitoring and rollback/pause owner assigned.",
    ])
    add_callout(doc, "Owner decisions still required", "Developer account type, public support email/domain, privacy and deletion URLs, target countries, age audience, pricing, final version name/code, production tester roster, and whether the beta label is removed before public release.", "warning")

    add_page_break(doc)
    doc.add_heading("31. Official Google Play sources", level=2)
    add_para(doc, "Policy and console requirements change. These official sources were checked on 9 August 2026; verify them again immediately before submission.")
    add_source(doc, "Get started with Play Console", "https://support.google.com/googleplay/android-developer/answer/6112435?hl=en-EN", "developer account setup, verification, and US$25 registration fee")
    add_source(doc, "Create and set up your app", "https://support.google.com/googleplay/android-developer/answer/9859152?hl=en", "app creation, store setup, and Android App Bundles")
    add_source(doc, "Prepare and roll out a release", "https://support.google.com/googleplay/android-developer/answer/9859348?hl=en", "AAB upload, release notes, review, and rollout")
    add_source(doc, "Use Play App Signing", "https://support.google.com/googleplay/android-developer/answer/9842756?hl=en", "upload key and Google-managed app-signing key")
    add_source(doc, "Target API level requirements", "https://developer.android.com/google/play/requirements/target-sdk", "API 36 requirement starting 31 August 2026")
    add_source(doc, "App testing requirements for new personal accounts", "https://support.google.com/googleplay/android-developer/answer/14151465?hl=en", "12 testers continuously opted in for 14 days")
    add_source(doc, "Data safety section", "https://support.google.com/googleplay/android-developer/answer/10787469?hl=en", "data collection/sharing and privacy disclosures")
    add_source(doc, "Health apps declaration", "https://support.google.com/googleplay/android-developer/answer/14738291?hl=en", "Medication and Treatment Management declaration")
    add_source(doc, "Health Content and Services policy", "https://support.google.com/googleplay/android-developer/answer/16679511?hl=en", "public privacy policy and health/medical requirements")
    add_source(doc, "Developer Program Policy", "https://support.google.com/googleplay/android-developer/answer/17190352?hl=en", "required non-medical-device disclaimer")
    add_source(doc, "Content rating requirements", "https://support.google.com/googleplay/android-developer/answer/9859655?hl=en", "IARC rating and target audience")
    add_source(doc, "Preview asset requirements", "https://support.google.com/googleplay/android-developer/answer/9866151?hl=en", "icon, feature graphic, and screenshot specifications")
    add_source(doc, "Permissions and sensitive APIs", "https://support.google.com/googleplay/android-developer/answer/16558241", "exact alarm permission policy")
    add_source(doc, "Android exact-alarm behavior", "https://developer.android.com/about/versions/14/changes/schedule-exact-alarms", "Android 14+ default denial and testing guidance")
    add_source(doc, "Registering Play package names", "https://support.google.com/googleplay/android-developer/answer/16984799?hl=en", "30 September 2026 package registration requirement")

    doc.add_heading("32. Project source references", level=2)
    add_bullets(doc, [
        "README.md and docs/README.md",
        "docs/ARCHITECTURE.md, docs/SECURITY.md, docs/TESTING.md, and docs/ANDROID_RELEASE.md",
        "package.json, capacitor.config.ts, android/variables.gradle, android/app/build.gradle, and AndroidManifest.xml",
        "Verified release artifacts under artifacts/ and bundled assistant audio under public/audio/assistant/",
    ])
    add_para(doc, "End of guide", italic=True, color=MID_GRAY)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "MedLoop AI - Project Details and Google Play Publishing Guide"
    doc.core_properties.subject = "Application architecture, release readiness, and Google Play upload SOP"
    doc.core_properties.author = "MedLoop AI Project"
    doc.core_properties.keywords = "MedLoop AI, Android, Google Play, release, privacy, health app"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
